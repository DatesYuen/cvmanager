import os
import uuid
import shutil
import io
import mimetypes
import zipfile
import difflib
from datetime import datetime
from typing import List, Optional
from pathlib import Path
from urllib.parse import quote
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from fastapi.responses import FileResponse, HTMLResponse, StreamingResponse
from pydantic import BaseModel
from sqlalchemy.orm import Session
from docx import Document

from backend.database import get_db
from backend.models import Attachment, AttachmentFolder, Person, User
from backend.schemas.common import AttachmentOut
from backend.services.auth_service import get_current_user
from backend.config import ATTACHMENT_UPLOAD_DIR
from backend.services.review_service import ENTITY_MODEL_MAP
from backend.services.export_service import SHOWCASE_ENTITY_LABELS

router = APIRouter(prefix="/api/attachments", tags=["attachments"])


class PlaceholderAttachmentRequest(BaseModel):
    entity_type: str
    entity_id: int
    display_name: str = ""


class AttachmentFolderCreate(BaseModel):
    person_id: int
    name: str
    parent_id: Optional[int] = None


class AttachmentExportRequest(BaseModel):
    person_id: int
    attachment_ids: List[int] = []


@router.get("/", response_model=List[AttachmentOut])
def list_attachments(entity_type: str, entity_id: int,
                     db: Session = Depends(get_db),
                     user: User = Depends(get_current_user)):
    return db.query(Attachment).filter(
        Attachment.entity_type == entity_type,
        Attachment.entity_id == entity_id
    ).all()


@router.get("/person/{person_id}")
def list_person_attachments(person_id: int,
                            db: Session = Depends(get_db),
                            user: User = Depends(get_current_user)):
    person = db.query(Person).get(person_id)
    if not person:
        raise HTTPException(status_code=404, detail="Person not found")

    folders = db.query(AttachmentFolder).filter(AttachmentFolder.person_id == person_id).all()
    folder_map = {folder.id: folder for folder in folders}
    items = []

    for attachment in db.query(Attachment).all():
        context = _attachment_context(db, attachment, folder_map)
        if not context or context["person_id"] != person_id:
            continue
        items.append({
            "id": attachment.id,
            "entity_type": attachment.entity_type,
            "entity_id": attachment.entity_id,
            "folder_id": attachment.folder_id,
            "original_filename": attachment.original_filename,
            "uploaded_at": attachment.uploaded_at,
            "category": context["category"],
            "item_name": context["item_name"],
            "hierarchy": context["hierarchy"],
            "is_placeholder": attachment.original_filename.startswith("【暂无】"),
        })

    items.sort(key=lambda item: (item["category"], item["item_name"], item["uploaded_at"]))
    return {
        "folders": [
            {
                "id": folder.id,
                "person_id": folder.person_id,
                "name": folder.name,
                "parent_id": folder.parent_id,
                "hierarchy": _folder_hierarchy(folder, folder_map),
                "created_at": folder.created_at,
            }
            for folder in folders
        ],
        "attachments": items,
    }


@router.post("/folders")
def create_attachment_folder(req: AttachmentFolderCreate,
                             db: Session = Depends(get_db),
                             user: User = Depends(get_current_user)):
    if not req.name.strip():
        raise HTTPException(status_code=400, detail="Folder name is required")
    person = db.query(Person).get(req.person_id)
    if not person:
        raise HTTPException(status_code=404, detail="Person not found")
    if req.parent_id:
        parent = db.query(AttachmentFolder).get(req.parent_id)
        if not parent or parent.person_id != req.person_id:
            raise HTTPException(status_code=400, detail="Invalid parent folder")

    folder = AttachmentFolder(
        person_id=req.person_id,
        name=req.name.strip(),
        parent_id=req.parent_id,
        created_by=user.id,
    )
    db.add(folder)
    db.commit()
    db.refresh(folder)
    return folder


@router.delete("/folders/{folder_id}")
def delete_attachment_folder(folder_id: int,
                             db: Session = Depends(get_db),
                             user: User = Depends(get_current_user)):
    folder = db.query(AttachmentFolder).get(folder_id)
    if not folder:
        raise HTTPException(status_code=404, detail="Folder not found")
    folder_ids = _collect_folder_descendant_ids(db, folder.id)
    attachments = db.query(Attachment).filter(Attachment.folder_id.in_(folder_ids)).all()
    for attachment in attachments:
        if attachment.file_path and os.path.exists(attachment.file_path):
            os.remove(attachment.file_path)
        db.delete(attachment)

    folder_rows = db.query(AttachmentFolder).filter(AttachmentFolder.id.in_(folder_ids)).all()
    folder_by_id = {item.id: item for item in folder_rows}
    for descendant_id in reversed(folder_ids):
        folder_to_delete = folder_by_id.get(descendant_id)
        if folder_to_delete:
            db.delete(folder_to_delete)
    db.commit()
    return {"deleted_folders": len(folder_ids), "deleted_attachments": len(attachments)}


@router.post("/folders/{folder_id}/upload", response_model=AttachmentOut)
def upload_folder_attachment(folder_id: int,
                             file: UploadFile = File(...),
                             db: Session = Depends(get_db),
                             user: User = Depends(get_current_user)):
    folder = db.query(AttachmentFolder).get(folder_id)
    if not folder:
        raise HTTPException(status_code=404, detail="Folder not found")
    original_filename = Path(file.filename or "attachment").name
    if not original_filename:
        raise HTTPException(status_code=400, detail="Filename is empty")

    filename = f"folder_{folder_id}_{uuid.uuid4().hex[:8]}_{original_filename}"
    file_path = ATTACHMENT_UPLOAD_DIR / filename
    with open(file_path, "wb") as f:
        shutil.copyfileobj(file.file, f)

    attachment = Attachment(
        entity_type="attachment_folders",
        entity_id=folder.id,
        folder_id=folder.id,
        file_path=str(file_path),
        original_filename=original_filename,
        uploaded_by=user.id,
    )
    db.add(attachment)
    db.commit()
    db.refresh(attachment)
    return attachment


@router.get("/match")
def match_attachment_targets(person_id: int, filename: str,
                             db: Session = Depends(get_db),
                             user: User = Depends(get_current_user)):
    person = db.query(Person).get(person_id)
    if not person:
        raise HTTPException(status_code=404, detail="Person not found")

    query = _normalize_match_text(Path(filename or "").stem)
    candidates = []
    for entity_type, Model in ENTITY_MODEL_MAP.items():
        for item in db.query(Model).filter(Model.person_id == person_id).all():
            item_name = _entity_display_name(entity_type, item)
            score = _match_score(query, _normalize_match_text(item_name))
            if score <= 0:
                continue
            candidates.append({
                "entity_type": entity_type,
                "entity_id": item.id,
                "category": SHOWCASE_ENTITY_LABELS.get(entity_type, entity_type),
                "item_name": item_name,
                "score": round(score, 4),
            })
    candidates.sort(key=lambda item: item["score"], reverse=True)
    return candidates[:30]


@router.post("/person-upload", response_model=AttachmentOut)
def upload_person_attachment(person_id: int = Form(...),
                             target_entity_type: Optional[str] = Form(None),
                             target_entity_id: Optional[int] = Form(None),
                             file: UploadFile = File(...),
                             db: Session = Depends(get_db),
                             user: User = Depends(get_current_user)):
    person = db.query(Person).get(person_id)
    if not person:
        raise HTTPException(status_code=404, detail="Person not found")

    if target_entity_type and target_entity_id:
        Model = ENTITY_MODEL_MAP.get(target_entity_type)
        if not Model:
            raise HTTPException(status_code=400, detail="Unsupported entity type")
        item = db.query(Model).get(target_entity_id)
        if not item or item.person_id != person_id:
            raise HTTPException(status_code=404, detail="Entity not found")
        return _save_entity_attachment(db, user.id, target_entity_type, target_entity_id, file, item)

    incoming_filename = Path(file.filename or "attachment").name
    if not incoming_filename:
        raise HTTPException(status_code=400, detail="Filename is empty")
    filename = f"person_{person_id}_{uuid.uuid4().hex[:8]}_{incoming_filename}"
    file_path = ATTACHMENT_UPLOAD_DIR / filename
    with open(file_path, "wb") as f:
        shutil.copyfileobj(file.file, f)
    attachment = Attachment(
        entity_type="person_uploads",
        entity_id=person_id,
        file_path=str(file_path),
        original_filename=incoming_filename,
        uploaded_by=user.id,
    )
    db.add(attachment)
    db.commit()
    db.refresh(attachment)
    return attachment


@router.post("/export")
def export_person_attachments(req: AttachmentExportRequest,
                              db: Session = Depends(get_db),
                              user: User = Depends(get_current_user)):
    person = db.query(Person).get(req.person_id)
    if not person:
        raise HTTPException(status_code=404, detail="Person not found")
    folder_map = {
        folder.id: folder
        for folder in db.query(AttachmentFolder).filter(AttachmentFolder.person_id == req.person_id).all()
    }
    q = db.query(Attachment)
    if req.attachment_ids:
        q = q.filter(Attachment.id.in_(req.attachment_ids))

    output = io.BytesIO()
    written = set()
    with zipfile.ZipFile(output, mode="w", compression=zipfile.ZIP_DEFLATED) as archive:
        count = 0
        for attachment in q.all():
            context = _attachment_context(db, attachment, folder_map)
            if not context or context["person_id"] != req.person_id:
                continue
            if not os.path.exists(attachment.file_path):
                continue
            zip_name = _dedupe_zip_name(
                f"{_safe_zip_path(context['hierarchy'])}/{_safe_filename_part(attachment.original_filename)}",
                written,
            )
            archive.write(attachment.file_path, zip_name)
            count += 1
        if count == 0:
            archive.writestr("README.txt", "没有可导出的附件。")
    output.seek(0)
    filename = f"{person.name or 'person'}_attachments_{datetime.now().strftime('%Y-%m-%d')}.zip"
    return StreamingResponse(
        output,
        media_type="application/zip",
        headers={
            "Content-Disposition": f"attachment; filename=attachments.zip; filename*=UTF-8''{quote(filename)}"
        },
    )


@router.get("/status")
def attachment_status(entity_type: str, entity_ids: str,
                      db: Session = Depends(get_db),
                      user: User = Depends(get_current_user)):
    ids = [_to_int(value) for value in entity_ids.split(",")]
    ids = [value for value in ids if value is not None]
    if not ids:
        return {}

    rows = db.query(Attachment).filter(
        Attachment.entity_type == entity_type,
        Attachment.entity_id.in_(ids),
    ).all()
    status = {
        str(entity_id): {"count": 0, "has_placeholder": False, "has_file": False}
        for entity_id in ids
    }
    for attachment in rows:
        item = status.setdefault(str(attachment.entity_id), {
            "count": 0,
            "has_placeholder": False,
            "has_file": False,
        })
        item["count"] += 1
        if attachment.original_filename.startswith("【暂无】"):
            item["has_placeholder"] = True
        else:
            item["has_file"] = True
    return status


@router.post("/upload", response_model=AttachmentOut)
def upload_attachment(entity_type: str, entity_id: int,
                      file: UploadFile = File(...),
                      db: Session = Depends(get_db),
                      user: User = Depends(get_current_user)):
    Model = ENTITY_MODEL_MAP.get(entity_type)
    if not Model:
        raise HTTPException(status_code=400, detail="Unsupported entity type")
    item = db.query(Model).get(entity_id)
    if not item:
        raise HTTPException(status_code=404, detail="Entity not found")
    return _save_entity_attachment(db, user.id, entity_type, entity_id, file, item)


@router.post("/placeholder", response_model=AttachmentOut)
def create_placeholder_attachment(req: PlaceholderAttachmentRequest,
                                  db: Session = Depends(get_db),
                                  user: User = Depends(get_current_user)):
    Model = ENTITY_MODEL_MAP.get(req.entity_type)
    if not Model:
        raise HTTPException(status_code=400, detail="Unsupported entity type")
    item = db.query(Model).get(req.entity_id)
    if not item:
        raise HTTPException(status_code=404, detail="Entity not found")

    display_name = req.display_name.strip() or _entity_display_name(req.entity_type, item)
    safe_display = _safe_filename_part(display_name)
    original_filename = f"【暂无】{safe_display}.txt"

    existing = db.query(Attachment).filter(
        Attachment.entity_type == req.entity_type,
        Attachment.entity_id == req.entity_id,
        Attachment.original_filename == original_filename,
    ).first()
    if existing:
        return existing

    disk_name = f"{req.entity_type}_{req.entity_id}_{uuid.uuid4().hex[:8]}_{original_filename}"
    file_path = ATTACHMENT_UPLOAD_DIR / disk_name
    file_path.write_text(f"{display_name}\n暂无附件。\n", encoding="utf-8")

    attachment = Attachment(
        entity_type=req.entity_type,
        entity_id=req.entity_id,
        file_path=str(file_path),
        original_filename=original_filename,
        uploaded_by=user.id,
    )
    db.add(attachment)
    db.commit()
    db.refresh(attachment)
    return attachment


@router.delete("/placeholder")
def delete_placeholder_attachment(entity_type: str, entity_id: int,
                                  db: Session = Depends(get_db),
                                  user: User = Depends(get_current_user)):
    placeholders = db.query(Attachment).filter(
        Attachment.entity_type == entity_type,
        Attachment.entity_id == entity_id,
    ).all()
    deleted = 0
    for attachment in placeholders:
        if not attachment.original_filename.startswith("【暂无】"):
            continue
        if os.path.exists(attachment.file_path):
            os.remove(attachment.file_path)
        db.delete(attachment)
        deleted += 1
    db.commit()
    return {"deleted": deleted}


@router.get("/preview/{attachment_id}")
def preview_attachment(attachment_id: int, db: Session = Depends(get_db),
                       user: User = Depends(get_current_user)):
    att = db.query(Attachment).get(attachment_id)
    if not att:
        raise HTTPException(status_code=404, detail="Attachment not found")
    if not os.path.exists(att.file_path):
        raise HTTPException(status_code=404, detail="File not found on disk")

    suffix = Path(att.original_filename).suffix.lower()
    if suffix == ".docx":
        return HTMLResponse(_docx_preview_html(att.file_path, att.original_filename))

    media_type = mimetypes.guess_type(att.original_filename)[0] or "application/octet-stream"
    if suffix in {".txt", ".md", ".csv", ".log"}:
        media_type = "text/plain; charset=utf-8"
    return FileResponse(
        att.file_path,
        media_type=media_type,
        filename=att.original_filename,
        content_disposition_type="inline",
    )


@router.get("/download/{attachment_id}")
def download_attachment(attachment_id: int, db: Session = Depends(get_db),
                        user: User = Depends(get_current_user)):
    att = db.query(Attachment).get(attachment_id)
    if not att:
        raise HTTPException(status_code=404, detail="Attachment not found")
    if not os.path.exists(att.file_path):
        raise HTTPException(status_code=404, detail="File not found on disk")
    return FileResponse(att.file_path, filename=att.original_filename)


@router.delete("/{attachment_id}")
def delete_attachment(attachment_id: int, db: Session = Depends(get_db),
                      user: User = Depends(get_current_user)):
    att = db.query(Attachment).get(attachment_id)
    if not att:
        raise HTTPException(status_code=404, detail="Attachment not found")
    if os.path.exists(att.file_path):
        os.remove(att.file_path)
    db.delete(att)
    db.commit()
    return {"msg": "Attachment deleted"}


def _entity_display_name(entity_type: str, item) -> str:
    for field in ("title", "name", "project_name", "award_name", "patent_name", "issue_name", "registration_number"):
        value = getattr(item, field, "")
        if value:
            return str(value)
    return f"{entity_type}_{item.id}"


def _safe_filename_part(value: str) -> str:
    cleaned = "".join("_" if ch in '<>:"/\\|?*\r\n\t' else ch for ch in str(value or "")).strip(" .")
    return cleaned[:120] or "未命名"


def _to_int(value: str) -> int | None:
    try:
        return int(value)
    except (TypeError, ValueError):
        return None


def _collect_folder_descendant_ids(db: Session, folder_id: int) -> list[int]:
    all_ids = [folder_id]
    cursor = 0
    while cursor < len(all_ids):
        current_id = all_ids[cursor]
        cursor += 1
        child_ids = [
            row[0]
            for row in db.query(AttachmentFolder.id).filter(AttachmentFolder.parent_id == current_id).all()
        ]
        for child_id in child_ids:
            if child_id not in all_ids:
                all_ids.append(child_id)
    return all_ids


def _attachment_context(db: Session, attachment: Attachment, folder_map: dict[int, AttachmentFolder]) -> dict | None:
    if attachment.entity_type == "person_uploads":
        person = db.query(Person).get(attachment.entity_id)
        if not person:
            return None
        return {
            "person_id": person.id,
            "category": "仅上传",
            "item_name": person.name,
            "hierarchy": "仅上传",
        }

    if attachment.entity_type == "attachment_folders":
        folder = folder_map.get(attachment.folder_id or attachment.entity_id)
        if not folder:
            folder = db.query(AttachmentFolder).get(attachment.folder_id or attachment.entity_id)
        if not folder:
            return None
        hierarchy = _folder_hierarchy(folder, folder_map)
        return {
            "person_id": folder.person_id,
            "category": "文件夹",
            "item_name": folder.name,
            "hierarchy": hierarchy,
        }

    Model = ENTITY_MODEL_MAP.get(attachment.entity_type)
    if not Model:
        return None
    item = db.query(Model).get(attachment.entity_id)
    if not item:
        return None
    category = SHOWCASE_ENTITY_LABELS.get(attachment.entity_type, attachment.entity_type)
    item_name = _entity_display_name(attachment.entity_type, item)
    return {
        "person_id": item.person_id,
        "category": category,
        "item_name": item_name,
        "hierarchy": f"{category}/{item_name}",
    }


def _folder_hierarchy(folder: AttachmentFolder, folder_map: dict[int, AttachmentFolder]) -> str:
    names = [folder.name]
    current = folder
    seen = {folder.id}
    while current.parent_id and current.parent_id not in seen:
        parent = folder_map.get(current.parent_id)
        if not parent:
            break
        names.append(parent.name)
        seen.add(parent.id)
        current = parent
    return "/".join(reversed(names))


def _dedupe_zip_name(name: str, seen: set[str]) -> str:
    candidate = name
    counter = 2
    path = Path(name)
    while candidate in seen:
        candidate = str(path.with_name(f"{path.stem}_{counter}{path.suffix}")).replace("\\", "/")
        counter += 1
    seen.add(candidate)
    return candidate


def _safe_zip_path(value: str) -> str:
    return "/".join(_safe_filename_part(part) for part in str(value or "").split("/") if part)


def _docx_preview_html(file_path: str, filename: str) -> str:
    doc = Document(file_path)
    body = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            body.append(f"<p>{_html_escape(text)}</p>")
    for table in doc.tables:
        body.append("<table>")
        for row in table.rows:
            body.append("<tr>")
            for cell in row.cells:
                body.append(f"<td>{_html_escape(cell.text.strip())}</td>")
            body.append("</tr>")
        body.append("</table>")
    return f"""
<!doctype html>
<html>
<head>
  <meta charset="utf-8" />
  <title>{_html_escape(filename)}</title>
  <style>
    body {{ font-family: Arial, "Microsoft YaHei", sans-serif; line-height: 1.7; padding: 24px; color: #303133; }}
    table {{ border-collapse: collapse; margin: 12px 0; width: 100%; }}
    td {{ border: 1px solid #dcdfe6; padding: 6px 8px; vertical-align: top; }}
  </style>
</head>
<body>
  <h2>{_html_escape(filename)}</h2>
  {''.join(body) or '<p>该 DOCX 未读取到可预览文本。</p>'}
</body>
</html>
"""


def _html_escape(value: str) -> str:
    return (
        str(value or "")
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


def _save_entity_attachment(db: Session, user_id: int, entity_type: str, entity_id: int,
                            file: UploadFile, item) -> Attachment:
    incoming_filename = Path(file.filename or "attachment").name
    if not incoming_filename:
        raise HTTPException(status_code=400, detail="Filename is empty")
    suffix = Path(incoming_filename).suffix
    original_filename = f"{_safe_filename_part(_entity_display_name(entity_type, item))}{suffix}"
    filename = f"{entity_type}_{entity_id}_{uuid.uuid4().hex[:8]}_{original_filename}"
    file_path = ATTACHMENT_UPLOAD_DIR / filename
    with open(file_path, "wb") as f:
        shutil.copyfileobj(file.file, f)

    attachment = Attachment(
        entity_type=entity_type,
        entity_id=entity_id,
        folder_id=None,
        file_path=str(file_path),
        original_filename=original_filename,
        uploaded_by=user_id,
    )
    db.add(attachment)
    db.commit()
    db.refresh(attachment)
    return attachment


def _normalize_match_text(value: str) -> str:
    text = str(value or "").lower()
    for ch in "_-—–·,，.。()（）[]【】":
        text = text.replace(ch, " ")
    return " ".join(text.split())


def _match_score(query: str, candidate: str) -> float:
    if not query or not candidate:
        return 0
    if query == candidate:
        return 1
    if query in candidate or candidate in query:
        shorter = min(len(query), len(candidate))
        longer = max(len(query), len(candidate))
        return 0.72 + 0.28 * (shorter / longer)
    score = difflib.SequenceMatcher(None, query, candidate).ratio()
    return score if score >= 0.35 else 0

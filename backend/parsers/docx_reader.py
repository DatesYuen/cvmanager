"""DOCX reader: extract paragraphs and tables from Word documents."""
from typing import List, Dict, Any
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def read_docx(file_path: str) -> Dict[str, Any]:
    """Read a DOCX file and return structured content.

    Returns:
        {
            "paragraphs": [{"text": str, "style": str, "bold": bool, "level": int}],
            "tables": [{"rows": [[str, ...], ...]}],
            "full_text": str
        }
    """
    doc = Document(file_path)
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style_name = para.style.name if para.style else ""
        is_bold = any(run.bold for run in para.runs if run.bold is not None)
        level = 0
        if "Heading" in style_name:
            try:
                level = int(style_name.split()[-1])
            except (ValueError, IndexError):
                level = 1
        paragraphs.append({
            "text": text,
            "style": style_name,
            "bold": is_bold,
            "level": level,
        })

    tables = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        tables.append({"rows": rows})

    full_text = "\n".join(p["text"] for p in paragraphs)
    return {
        "paragraphs": paragraphs,
        "tables": tables,
        "full_text": full_text,
    }

"""Export service: build filtered data for export."""
import io
from datetime import datetime
from pathlib import Path

from sqlalchemy import Boolean
from sqlalchemy.orm import Session
from docx import Document
from docx.oxml.ns import qn
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfbase import pdfmetrics
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from backend.models import (
    Person, Profile, Education, WorkExperience, Attachment,
    Paper, Project, Award, Patent, SoftwareCopyright,
    StudentAward, Conference, SpecialIssue, AcademicRole,
    AcademicReport, TeachingPlatform, IndustryStandard,
)
from backend.schemas.common import ExportRequest
from backend.services.review_service import ENTITY_MODEL_MAP


def build_export_data(db: Session, req: ExportRequest):
    Model = ENTITY_MODEL_MAP.get(req.entity_type)
    if not Model:
        return [], []

    q = db.query(Model)
    if req.person_id:
        q = q.filter(Model.person_id == req.person_id)

    # Apply filters
    for f in req.filters:
        col = getattr(Model, f.field, None)
        if col is None:
            continue
        value = _normalize_filter_value(col, f.value)
        if f.op == "eq":
            q = q.filter(col == value)
        elif f.op == "ne":
            q = q.filter(col != value)
        elif f.op == "contains":
            q = q.filter(col.contains(str(value)))
        elif f.op == "gt":
            q = q.filter(col > value)
        elif f.op == "lt":
            q = q.filter(col < value)
        elif f.op == "gte":
            q = q.filter(col >= value)
        elif f.op == "lte":
            q = q.filter(col <= value)

    items = q.all()

    # Build column list
    all_columns = [c.name for c in Model.__table__.columns if c.name != "id"]
    columns = req.fields if req.fields else all_columns

    data = []
    for item in items:
        row = {}
        for c in columns:
            row[c] = getattr(item, c, "")
        data.append(row)

    return data, columns


def _normalize_filter_value(column, value):
    try:
        python_type = column.property.columns[0].type.python_type
    except Exception:  # noqa: BLE001
        python_type = None

    if python_type is bool:
        if isinstance(value, bool):
            return value
        lowered = str(value).strip().lower()
        return lowered in {"true", "1", "yes", "y", "是"}

    return value


SHOWCASE_ENTITY_MODELS = {
    "papers": Paper,
    "projects": Project,
    "awards": Award,
    "patents": Patent,
    "software_copyrights": SoftwareCopyright,
    "student_awards": StudentAward,
    "conferences": Conference,
    "special_issues": SpecialIssue,
    "academic_roles": AcademicRole,
    "academic_reports": AcademicReport,
    "teaching_platforms": TeachingPlatform,
    "industry_standards": IndustryStandard,
}

SHOWCASE_ENTITY_LABELS = {
    "papers": "论文",
    "projects": "项目",
    "awards": "获奖",
    "patents": "专利",
    "software_copyrights": "软著",
    "student_awards": "指导学生获奖",
    "conferences": "承办会议",
    "special_issues": "承办特刊",
    "academic_roles": "学术兼职",
    "academic_reports": "学术报告",
    "teaching_platforms": "教学平台建设",
    "industry_standards": "行业标准",
}


def build_docx_export(db: Session, req: ExportRequest) -> tuple[io.BytesIO, str]:
    return build_resume_docx_export(db, req.person_id)


def build_resume_docx_export(db: Session, person_id: int | None) -> tuple[io.BytesIO, str]:
    if not person_id:
        raise ValueError("DOCX导出需要先选择一个人员")

    person = db.query(Person).get(person_id)
    if not person:
        raise ValueError("Person not found")

    doc = Document()
    _apply_songti_font(doc)
    doc.add_heading(person.name, level=0)
    if person.name_en:
        doc.add_paragraph(person.name_en)

    _append_profile_section(doc, db, person.id)

    for entity_type, Model in SHOWCASE_ENTITY_MODELS.items():
        items = db.query(Model).filter(
            Model.person_id == person.id,
            Model.review_status == "approved",
        ).all()
        if not items:
            continue
        doc.add_heading(SHOWCASE_ENTITY_LABELS[entity_type], level=2)
        for idx, item in enumerate(items, start=1):
            doc.add_paragraph(f"{idx}. {_format_item(entity_type, item)}")
            attachment_line = _format_attachments(db, entity_type, item.id)
            if attachment_line:
                doc.add_paragraph(attachment_line)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output, _build_person_filename(person.name, ".docx")


def build_entity_docx_export(db: Session, req: ExportRequest) -> tuple[io.BytesIO, str]:
    if not req.entity_type:
        raise ValueError("条目DOCX导出需要选择数据类型")
    Model = ENTITY_MODEL_MAP.get(req.entity_type)
    if not Model:
        raise ValueError("Unsupported entity type")

    person = db.query(Person).get(req.person_id) if req.person_id else None
    label = SHOWCASE_ENTITY_LABELS.get(req.entity_type, req.entity_type)
    doc = Document()
    _apply_songti_font(doc)
    doc.add_heading(label, level=0)
    if person:
        doc.add_paragraph(f"人员：{person.name}")
        if person.name_en:
            doc.add_paragraph(person.name_en)

    q = db.query(Model)
    if req.person_id:
        q = q.filter(Model.person_id == req.person_id)

    for f in req.filters:
        col = getattr(Model, f.field, None)
        if col is None:
            continue
        value = _normalize_filter_value(col, f.value)
        if f.op == "eq":
            q = q.filter(col == value)
        elif f.op == "ne":
            q = q.filter(col != value)
        elif f.op == "contains":
            q = q.filter(col.contains(str(value)))
        elif f.op == "gt":
            q = q.filter(col > value)
        elif f.op == "lt":
            q = q.filter(col < value)
        elif f.op == "gte":
            q = q.filter(col >= value)
        elif f.op == "lte":
            q = q.filter(col <= value)

    items = q.all()
    for idx, item in enumerate(items, start=1):
        doc.add_paragraph(f"{idx}. {_format_item(req.entity_type, item)}")

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    if person:
        filename = _build_person_filename(person.name, ".docx")
    else:
        filename = f"{req.entity_type}_{datetime.now().strftime('%Y-%m-%d')}.docx"
    return output, filename


def build_resume_pdf_export(db: Session, person_id: int | None) -> tuple[io.BytesIO, str]:
    if not person_id:
        raise ValueError("PDF导出需要先选择一个人员")

    person = db.query(Person).get(person_id)
    if not person:
        raise ValueError("Person not found")

    output = io.BytesIO()
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "CvTitle",
        parent=styles["Title"],
        fontName="STSong-Light",
        fontSize=18,
        leading=24,
    )
    heading_style = ParagraphStyle(
        "CvHeading",
        parent=styles["Heading2"],
        fontName="STSong-Light",
        fontSize=14,
        leading=20,
        spaceBefore=10,
        spaceAfter=6,
    )
    body_style = ParagraphStyle(
        "CvBody",
        parent=styles["BodyText"],
        fontName="STSong-Light",
        fontSize=11,
        leading=16,
    )

    story = [Paragraph(person.name, title_style)]
    if person.name_en:
        story.append(Paragraph(person.name_en, body_style))
        story.append(Spacer(1, 4))

    profile = db.query(Profile).filter(Profile.person_id == person.id).first()
    educations = db.query(Education).filter(Education.person_id == person.id).all()
    works = db.query(WorkExperience).filter(WorkExperience.person_id == person.id).all()

    if profile:
        story.append(Paragraph("个人简介", heading_style))
        for line in filter(None, [
            profile.introduction,
            f"联系电话：{profile.phone}" if profile.phone else "",
            f"电子邮箱：{profile.email}" if profile.email else "",
            f"联系地址：{profile.address}" if profile.address else "",
        ]):
            story.append(Paragraph(line.replace("\n", "<br/>"), body_style))

    if educations:
        story.append(Paragraph("学习经历", heading_style))
        for idx, item in enumerate(educations, start=1):
            story.append(Paragraph(
                f"{idx}. {item.start_date} - {item.end_date}，{item.school}，{item.major}，{item.degree}",
                body_style
            ))

    if works:
        story.append(Paragraph("工作经历", heading_style))
        for idx, item in enumerate(works, start=1):
            story.append(Paragraph(
                f"{idx}. {item.start_date} - {item.end_date}，{item.organization}，{item.position}",
                body_style
            ))

    for entity_type, Model in SHOWCASE_ENTITY_MODELS.items():
        items = db.query(Model).filter(
            Model.person_id == person.id,
            Model.review_status == "approved",
        ).all()
        if not items:
            continue
        story.append(Paragraph(SHOWCASE_ENTITY_LABELS[entity_type], heading_style))
        for idx, item in enumerate(items, start=1):
            story.append(Paragraph(f"{idx}. {_format_item(entity_type, item)}", body_style))
            attachment_line = _format_attachments(db, entity_type, item.id)
            if attachment_line:
                story.append(Paragraph(attachment_line, body_style))

    doc = SimpleDocTemplate(
        output,
        pagesize=A4,
        leftMargin=20 * mm,
        rightMargin=20 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
    )
    doc.build(story)
    output.seek(0)
    return output, _build_person_filename(person.name, ".pdf")


def _append_profile_section(doc: Document, db: Session, person_id: int) -> None:
    profile = db.query(Profile).filter(Profile.person_id == person_id).first()
    educations = db.query(Education).filter(Education.person_id == person_id).all()
    works = db.query(WorkExperience).filter(WorkExperience.person_id == person_id).all()

    if profile:
        doc.add_heading("个人简介", level=2)
        if profile.introduction:
            doc.add_paragraph(profile.introduction)
        if profile.phone:
            doc.add_paragraph(f"联系电话：{profile.phone}")
        if profile.email:
            doc.add_paragraph(f"电子邮箱：{profile.email}")
        if profile.address:
            doc.add_paragraph(f"联系地址：{profile.address}")

    if educations:
        doc.add_heading("学习经历", level=2)
        for idx, item in enumerate(educations, start=1):
            doc.add_paragraph(
                f"{idx}. {item.start_date} - {item.end_date}，{item.school}，{item.major}，{item.degree}"
            )

    if works:
        doc.add_heading("工作经历", level=2)
        for idx, item in enumerate(works, start=1):
            doc.add_paragraph(
                f"{idx}. {item.start_date} - {item.end_date}，{item.organization}，{item.position}"
            )


def _format_item(entity_type: str, item) -> str:
    if entity_type == "papers":
        return _format_paper(item)
    if entity_type == "patents":
        return _format_patent(item)
    if entity_type == "projects":
        return "；".join(filter(None, [
            item.project_type,
            item.name,
            f"项目编号：{item.project_number}" if item.project_number else "",
            f"{item.start_date} - {item.end_date}" if item.start_date or item.end_date else "",
            item.role,
            f"{item.amount}" if item.amount else "",
        ]))
    if entity_type == "awards":
        return "；".join(filter(None, [item.award_name, item.project_name, item.participants, item.awarding_body]))
    if entity_type == "software_copyrights":
        return "；".join(filter(None, [item.applicant, item.name, item.registration_date, item.registration_number]))
    if entity_type == "student_awards":
        return "；".join(filter(None, [item.award_name, item.level, item.role, item.award_date]))
    if entity_type == "conferences":
        return "；".join(filter(None, [item.name, item.date, item.role, item.website]))
    if entity_type == "special_issues":
        return "；".join(filter(None, [item.issue_name, item.journal_name, item.date, item.role]))
    if entity_type == "academic_roles":
        return "；".join(filter(None, [item.title, item.start_date, item.end_date]))
    if entity_type == "academic_reports":
        return "；".join(filter(None, [item.name, item.report_type, item.date]))
    if entity_type == "teaching_platforms":
        return "；".join(filter(None, [item.name, item.issuing_body, item.approval_date, item.position]))
    if entity_type == "industry_standards":
        return "；".join(filter(None, [item.name, item.publish_date, item.role]))
    return str(item)


def _format_paper(item: Paper) -> str:
    # Follows the currently effective national bibliographic convention in spirit
    # (GB/T 7714-2015 as of 2026-04-17).
    authors = ", ".join(
        f"{author.name}{'*' if author.is_corresponding_author else ''}" for author in item.authors
    )
    parts = [
        authors,
        f"{item.title}[J]",
        item.journal,
        item.year,
    ]
    vol_issue = ""
    if item.volume and item.issue:
        vol_issue = f"{item.volume}({item.issue})"
    elif item.volume:
        vol_issue = item.volume
    elif item.issue:
        vol_issue = f"({item.issue})"
    if vol_issue:
        parts.append(vol_issue)
    if item.pages:
        parts.append(item.pages)
    if item.doi:
        parts.append(f"DOI: {item.doi}")
    return ", ".join(filter(None, parts))


def _format_patent(item: Patent) -> str:
    applicants = ", ".join(applicant.name for applicant in item.applicants)
    parts = [
        applicants,
        f"{item.patent_name}[P]",
        f"申请号：{item.application_number}" if item.application_number else "",
        f"授权号：{item.authorization_number}" if item.authorization_number else "",
        item.status,
    ]
    return "，".join(filter(None, parts))


def _format_attachments(db: Session, entity_type: str, entity_id: int) -> str:
    attachments = db.query(Attachment).filter(
        Attachment.entity_type == entity_type,
        Attachment.entity_id == entity_id,
    ).all()
    if not attachments:
        return ""
    names = "；".join(attachment.original_filename for attachment in attachments)
    return f"附件：{names}"


def _build_person_filename(person_name: str, suffix: str) -> str:
    safe_name = person_name or "person"
    date_str = datetime.now().strftime("%Y-%m-%d")
    return f"{safe_name}_{date_str}{suffix}"


def _apply_songti_font(doc: Document) -> None:
    styles = doc.styles
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        if style_name in styles:
            font = styles[style_name].font
            font.name = "宋体"
            if styles[style_name]._element.rPr is not None:
                styles[style_name]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
